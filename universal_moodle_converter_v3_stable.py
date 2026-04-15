# -*- coding: utf-8 -*-
"""
Универсальный конвертер Word -> Moodle XML
Версия 3.0 (Stable) — финальная стабильная версия

Поддерживаемые типы вопросов:
  - shortanswer  (краткий ответ)
  - multichoice  (множественный выбор)
  - cloze        (пропущенные слова)
  - matching     (сопоставление)

Особенности:
  - Извлечение изображений из docx (base64)
  - Конвертация формул LaTeX ($...$ → \\(...\\))
  - Обработка повреждённых маркеров вопросов (5 форматов)
"""

import os
import sys
import base64
import re
import unicodedata
import logging
import argparse
import itertools
from datetime import datetime
from typing import List, Dict, Set
from lxml import etree
from docx import Document
from docx.parts.image import ImagePart

try:
    import docxlatex
except ImportError:
    print("Внимание: docxlatex не установлен. Установите: pip install docxlatex")
    docxlatex = None


# ============================================================
# НАСТРОЙКИ ЛОГИРОВАНИЯ
# ============================================================
LOG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')
os.makedirs(LOG_DIR, exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(
            os.path.join(LOG_DIR, f'converter_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'),
            encoding='utf-8'
        ),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)


# ============================================================
# XML ШАБЛОНЫ ДЛЯ РАЗНЫХ ТИПОВ ВОПРОСОВ
# ============================================================

SHORTANSWER_TEMPLATE = '''<question type="shortanswer">
    <name><text/></name>
    <questiontext format="html"><text/></questiontext>
    <generalfeedback format="html"><text/></generalfeedback>
    <defaultgrade>1</defaultgrade>
    <penalty>0.3333333</penalty>
    <hidden>0</hidden>
    <idnumber/>
    <usecase>0</usecase>
    <correctfeedback format="html"><text/></correctfeedback>
    <partiallycorrectfeedback format="html"><text/></partiallycorrectfeedback>
    <incorrectfeedback format="html"><text/></incorrectfeedback>
</question>'''

MULTICHOICE_TEMPLATE = '''<question type="multichoice">
    <name><text/></name>
    <questiontext format="html"><text/></questiontext>
    <generalfeedback format="html"><text/></generalfeedback>
    <defaultgrade>1</defaultgrade>
    <penalty>0.3333333</penalty>
    <hidden>0</hidden>
    <idnumber/>
    <single>true</single>
    <shuffleanswers>true</shuffleanswers>
    <answernumbering>none</answernumbering>
    <showstandardinstruction>0</showstandardinstruction>
    <correctfeedback format="html"><text/></correctfeedback>
    <partiallycorrectfeedback format="html"><text/></partiallycorrectfeedback>
    <incorrectfeedback format="html"><text/></incorrectfeedback>
    <shownumcorrect/>
</question>'''

CLOZE_TEMPLATE = '''<question type="cloze">
    <name><text/></name>
    <questiontext format="html"><text/></questiontext>
    <generalfeedback format="html"/>
    <defaultgrade>1</defaultgrade>
    <penalty>0.3333333</penalty>
    <hidden>0</hidden>
    <idnumber/>
</question>'''

MATCHING_TEMPLATE = '''<question type="matching">
    <name><text/></name>
    <questiontext format="html"><text/></questiontext>
    <generalfeedback format="html"/>
    <defaultgrade>1</defaultgrade>
    <penalty>0.3333333</penalty>
    <hidden>0</hidden>
    <idnumber/>
    <subquestionsorten>false</subquestionsorten>
</question>'''

DDMATCH_TEMPLATE = '''<question type="ddmatch">
    <name><text/></name>
    <questiontext format="html"><text/></questiontext>
    <generalfeedback format="html"/>
    <defaultgrade>1</defaultgrade>
    <penalty>0.3333333</penalty>
    <hidden>0</hidden>
    <idnumber/>
    <shuffleanswers>true</shuffleanswers>
</question>'''

GAPSELECT_TEMPLATE = '''<question type="gapselect">
    <name><text/></name>
    <questiontext format="html"><text/></questiontext>
    <generalfeedback format="html"><text/></generalfeedback>
    <defaultgrade>1</defaultgrade>
    <penalty>0.3333333</penalty>
    <hidden>0</hidden>
    <idnumber/>
    <shuffleanswers>true</shuffleanswers>
    <correctfeedback format="html"><text/></correctfeedback>
    <partiallycorrectfeedback format="html"><text/></partiallycorrectfeedback>
    <incorrectfeedback format="html"><text/></incorrectfeedback>
    <shownumcorrect/>
</question>'''

NUMERICAL_TEMPLATE = '''<question type="numerical">
    <name><text/></name>
    <questiontext format="html"><text/></questiontext>
    <generalfeedback format="html"><text></text></generalfeedback>
    <defaultgrade>1</defaultgrade>
    <penalty>0.3333333</penalty>
    <hidden>0</hidden>
    <idnumber/>
    <answer fraction="100" format="moodle_auto_format"><text/><feedback format="html"><text></text></feedback><tolerance>0</tolerance></answer>
    <unitgradingtype>0</unitgradingtype>
    <unitpenalty>0.1000000</unitpenalty>
    <showunits>3</showunits>
    <unitsleft>0</unitsleft>
    <correctfeedback format="html"><text/></correctfeedback>
    <partiallycorrectfeedback format="html"><text/></partiallycorrectfeedback>
    <incorrectfeedback format="html"><text/></incorrectfeedback>
</question>'''


# ============================================================
# КЛАССЫ ДЛЯ ОБРАБОТКИ ИЗОБРАЖЕНИЙ И ФОРМУЛ
# ============================================================

class ImageProcessor:
    """Обработка изображений из Word документов.
    
    Извлекает все изображения из docx файла, кодирует в base64.
    Использует формат reference scripts: image1, image2, ...
    """
    
    def __init__(self):
        self.image_data: Dict[str, str] = {}
        self.image_by_para: Dict[int, List[str]] = {}
    
    def extract_images(self, docx_path: str) -> Dict[str, str]:
        """Извлекает все изображения из docx файла."""
        try:
            doc = Document(docx_path)
            self.image_data = {}
            self.image_by_para = {}
            n = 1
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_images = []
                rIds = re.findall(r'embed="([^"]+)"', paragraph._p.xml)
                
                for rId in rIds:
                    if rId in paragraph.part.rels:
                        rel = paragraph.part.rels[rId]
                        if isinstance(rel.target_part, ImagePart):
                            image_bytes = rel.target_part.blob
                            base64_str = base64.b64encode(image_bytes).decode('utf-8')
                            img_key = f'image{str(n)}'
                            self.image_data[img_key] = base64_str
                            para_images.append(img_key)
                            n += 1
                
                if para_images:
                    self.image_by_para[para_idx] = para_images
            
            logger.info(f"Извлечено изображений: {len(self.image_data)}")
            return self.image_data
            
        except Exception as e:
            logger.error(f"Ошибка при извлечении изображений: {e}")
            return {}
    
    @staticmethod
    def process_text(text: str, image_data: Dict[str, str]) -> str:
        """Обрабатывает текст: заменяет IMAGE#N-imageN маркеры.
        
        Маркеры формата: IMAGE#N-imageM или -N-imageM
        Преобразуются в _@@PLUGINFILE@@/imageM.png_IMAGE_base64_IMAGE_
        """
        def replace_image_marker(match):
            img_key = match.group(1)
            if img_key in image_data:
                return f'_@@PLUGINFILE@@/{img_key}.png_IMAGE_{image_data[img_key]}_IMAGE_'
            return match.group(0)
        
        text = re.sub(r'IMAGE#\d+-([\w]+)', replace_image_marker, text)
        text = re.sub(r'-(\d+)-([\w]+)', replace_image_marker, text)
        
        return text


def get_image(question_text):
    """Обрабатывает текст с IMAGE маркерами.
    
    Возвращает: [processed_text, [(image_name, base64_data), ...]]
    """
    images = []
    flag = False
    question = []
    current_img_name = ''
    current_img_data = ''
    
    if isinstance(question_text, list):
        quest_remodded = question_text[0].split('_')
    else:
        quest_remodded = question_text.split('_')
    
    for part in quest_remodded:
        if flag and part != 'IMAGE':
            current_img_data = part
        elif '@@PLUGINFILE@' in part:
            question.append('<p><img src="' + part + '"> </p>')
            current_img_name = part[part.find('/')+1:].replace('"', '')
        elif part == 'IMAGE':
            flag = not flag
            if flag:
                pass
            else:
                if current_img_name and current_img_data:
                    images.append((current_img_name, current_img_data))
                    current_img_name = ''
                    current_img_data = ''
        else:
            question.append('<p>' + part + '</p>')
    
    return [''.join(question), images]


class FormulaProcessor:
    """Обработка математических формул LaTeX."""
    
    @staticmethod
    def tex_to_latex(text: str) -> str:
        r"""Конвертирует $...$ в \(...\) для Moodle."""
        if '$' not in text:
            return text
        
        result = ''
        flag = False
        for char in text:
            if char == '$' and not flag:
                result += r'\('
                flag = True
            elif char == '$' and flag:
                result += r'\)'
                flag = False
            else:
                result += char
        return result
    
    @staticmethod
    def process_text(text: str, image_data: Dict[str, str]) -> str:
        """Обрабатывает текст: заменяет формулы и изображения."""
        text = FormulaProcessor.tex_to_latex(text)
        text = ImageProcessor.process_text(text, image_data)
        return text


# ============================================================
# МАППИНГ МАРКЕРОВ К ТИПАМ ВОПРОСОВ
# ============================================================

# Допустимые маркеры в формате {marker}V2:
VALID_MARKERS = {
    'multichoice_one',      # Один правильный ответ
    'multichoice_many',     # Несколько правильных, штраф -100% за неправильный
    'shortanswer_phrase',   # Запись текста (фразы)
    'numerical_partial',    # Цифры с partial scoring (100/50/0)
    'numerical_numcombo',   # Цифры в любом порядке (все перестановки = 100%)
    'matching',             # Соотношение (L/R пары)
    'match_123',            # Последовательность (фразы → номера позиций)
    'match',                # Соотношение (синоним matching)
    'ddmatch',              # Drag-and-drop matching
    'gapselect',            # Выпадающие списки по тексту
    'cloze',                # Выбор внутри текста вопроса
    'numerical',            # Числовой ответ (. и , как одно)
}

# Маппинг маркера → внутренний тип вопроса
MARKER_TO_QTYPE = {
    'multichoice_one':      'multichoice_one',
    'multichoice_many':     'multichoice_many',
    'shortanswer_phrase':   'shortanswer_phrase',
    'numerical_partial':    'numerical_partial',
    'numerical_numcombo':   'numerical_numcombo',
    'matching':             'matching',
    'match_123':            'match_123',
    'match':                'matching',        # match = matching
    'ddmatch':              'ddmatch',
    'gapselect':            'gapselect',
    'cloze':                'cloze',
    'numerical':            'numerical',
}


# ============================================================
# КЛАСС ДЛЯ ОПРЕДЕЛЕНИЯ ТИПА ВОПРОСА
# ============================================================

class QuestionTypeDetector:
    """Определение типа вопроса.
    
    Приоритет:
      1. Маркер из V2-строки ({marker}V2:...) — если есть
      2. Эвристика по содержимому — fallback
    """
    
    @staticmethod
    def detect(content: List[str], subject: str = '', marker: str = '') -> str:
        """Определяет тип вопроса.
        
        Args:
            content: список строк вопроса
            subject: код предмета ('ист', 'общ' и т.д.)
            marker: маркер из V2-строки ('' если нет)
        """
        # --- Приоритет 1: маркер ---
        if marker and marker in MARKER_TO_QTYPE:
            return MARKER_TO_QTYPE[marker]
        has_plus = any(item.startswith('+:') for item in content)
        has_plus_no_colon = any(item.startswith('+') and not item.startswith('+:') for item in content)
        has_minus = any(item.startswith('-:') or item.startswith('–:') for item in content)
        
        # Cloze: строгий формат Moodle {1:SHORTANSWER:answer} или ТЗ32-ТЗ35 в ИНФ
        has_cloze = any(
            re.match(r'^\{\d+:[A-Z]+:', item) for item in content
        )
        if not has_cloze:
            full_content = ' '.join(content)
            if re.search(r'[ТT](3[2-5])-\d+', full_content):
                has_plus_for_cloze = any(item.startswith('+:') for item in content)
                has_minus_for_cloze = any(item.startswith('-:') for item in content)
                if has_plus_for_cloze and has_minus_for_cloze:
                    has_cloze = True
        
        # Gapselect: (N) маркеры в тексте + строки с A) вариантами + ОТВЕТ:ABCD ключ
        has_gap_markers = any(re.search(r'\(\d+\)', item) for item in content)
        has_letter_options = any(re.match(r'^(\t| |)?\d*\.?\s*A\)', item) for item in content)
        has_gapselect_key = any(
            re.match(r'^ОТВЕТ:\s*[A-D]+$', item.strip()) or
            re.match(r'^Ответ:\s*[A-D\s]+$', item.strip()) or
            re.match(r'^Ответы\s*:\s*[A-D]+$', item.strip()) or
            re.match(r'^ответы\s*:\s*[A-D]+$', item.strip()) or
            re.match(r'^\+:\s*[A-D]+\s*$', item.strip())
            for item in content
        )
        
        has_gapselect = has_gap_markers and has_letter_options and has_gapselect_key
        
        # Numerical: один +: с числовым ответом, без -:
        has_numerical = False
        if has_plus and not has_minus and not has_gapselect:
            plus_items = [item.replace('+:', '').strip() for item in content if item.startswith('+:')]
            if len(plus_items) == 1 and re.match(r'^-?\d+(\.\d+)?$', plus_items[0]):
                full_content = ' '.join(content)
                if re.search(r'[ТT]36-\d+', full_content):
                    has_numerical = True
        
        # Matching/DDMATCH: L/R маркеры или стрелки
        has_matching_arrows = any('->' in item or '=>' in item for item in content)
        has_left_markers = any(re.match(r'^L\d+:', item) for item in content)
        has_right_markers = any(re.match(r'^R\d+:', item) for item in content)
        
        # Проверка на "Расположите фразы" - matching с нумерованными элементами
        has_numbered_items = any(re.match(r'^\d+:', item) for item in content)
        
        # Проверка инструкции на тип вопроса
        full_content = ' '.join(content)
        # ddmatch: инструкция содержит "тип" или "[грамматическая категория] + роль/норма"
        instruction_has_type = bool(re.search(r'\b(?:тип|типы)\b', full_content, re.IGNORECASE)) or \
                              bool(re.search(r'(?:синтаксическ|морфологическ|орфографическ|пунктуационн|фонетическ|лексическ|стилистическ)\w*\s+(?:роль|навык|норма)', full_content, re.IGNORECASE))
        instruction_has_headline = bool(re.search(r'\bзаголовок[ауоме]?\b', full_content, re.IGNORECASE))
        
        # Различаем matching и ddmatch по соотношению длин L и R и инструкции
        # ddmatch: R содержит типы/синонимы, инструкция содержит "тип"
        # matching: R содержит заголовки, инструкция содержит "заголовок"
        has_matching = False
        has_ddmatch = False
        
        if has_matching_arrows:
            has_matching = True
        elif has_numbered_items:
            # "Расположите фразы..." - matching с нумерованными элементами
            has_matching = True
        elif has_left_markers and has_right_markers:
            # L + R маркеры = matching (соотношение L/R)
            has_matching = True
            has_ddmatch = False
        
        # Определение типа по приоритету
        if has_cloze:
            return 'cloze'
        elif has_gapselect:
            return 'gapselect'
        elif has_numerical:
            return 'numerical'
        elif has_ddmatch:
            return 'ddmatch'
        elif has_matching:
            return 'matching'
        elif (has_plus or has_plus_no_colon) and has_minus:
            # Для ИСТ и ОБЩ используем shortanswer с перестановками (partial scoring)
            if subject in ('ист', 'общ'):
                return 'shortanswer'
            return 'multichoice'
        else:
            return 'shortanswer'


# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================================

def strip_service_markers(text: str) -> str:
    """Удаляет все сервисные маркеры (S:, I:, V1:, V2:, +:, -:) из начала текста."""
    while text.startswith(('S:', 'I:', 'V1:', 'V2:')):
        for prefix in ('V1:', 'V2:', 'S:', 'I:'):
            if text.startswith(prefix):
                text = text[len(prefix):].strip()
    return text


def remove_service_markers(text: str) -> str:
    """Удаляет все сервисные маркеры (S:, I:, V1:, V2:, +:, -:) из любого места в тексте."""
    import re
    text = re.sub(r'S:\s*', '', text)
    text = re.sub(r'I:\s*', '', text)
    text = re.sub(r'V1:\s*', '', text)
    text = re.sub(r'V2:\s*', '', text)
    text = re.sub(r'\+:\s*', '', text)
    text = re.sub(r'-:\s*', '', text)
    return text


def parse_answers_from_line(line: str) -> List[tuple]:
    """Парсит ответы из строки с разделителями ';' или ';' на конце.
    
    Примеры:
    - '+:либерализм;' -> [('либерализм', True)]
    - '-: монархия; -: унитарное государство; -: президентская республика; ' -> 
        [('монархия', False), ('унитарное государство', False), ('президентская республика', False)]
    """
    results = []
    line = line.strip()
    
    # Ищем все вхождения +: и -: в строке
    import re
    pattern = r'([+-]:\s*)([^;]+?)(?=\s*[+-]:|$)'
    matches = re.findall(pattern, line)
    
    for prefix, answer in matches:
        prefix = prefix.strip()
        is_right = prefix.startswith('+')
        answer = answer.strip().rstrip(';').strip()
        if answer:
            results.append((answer, is_right))
    
    # Если не нашли через regex (старая логика)
    if not results:
        if line.startswith('+:'):
            results.append((line.replace('+:', '').strip().rstrip(';').strip(), True))
        elif line.startswith('-:'):
            results.append((line.replace('-:', '').strip().rstrip(';').strip(), False))
    
    return results


# ============================================================
# КЛАСС ГЕНЕРАТОРА XML
# ============================================================

class XMLGenerator:
    """Генератор Moodle XML из структурированных данных."""
    
    def __init__(self):
        self.root = etree.Element('quiz')
        self.used_images: Set[str] = set()
        self.current_category = ""
        self.current_subcategory = ""
        self.image_data: Dict[str, str] = {}
        self.question_count = 0
    
    def set_image_data(self, image_data: Dict[str, str]):
        self.image_data = image_data
    
    def add_category(self, name: str):
        """Добавляет категорию вопросов."""
        if not name:
            return
        # Не добавляем дубликат категории
        if self.current_category == name:
            return
        path = f"$module$/top/{name}"
        cat_question = etree.SubElement(self.root, 'question', type='category')
        cat = etree.SubElement(cat_question, 'category')
        cat_text = etree.SubElement(cat, 'text')
        cat_text.text = path
        self.current_category = name
        # Сбрасываем подкатегорию при смене категории
        self.current_subcategory = ""
        logger.info(f"Добавлена категория: {name}")
    
    def add_subcategory(self, name: str):
        """Добавляет подкатегорию вопросов."""
        if not name:
            return
        # Не добавляем дубликат подкатегории
        if self.current_subcategory == name:
            return
        if self.current_category:
            path = f"$module$/top/{self.current_category}/{name}"
        else:
            path = f"$module$/top/{name}"
        cat_question = etree.SubElement(self.root, 'question', type='category')
        cat = etree.SubElement(cat_question, 'category')
        cat_text = etree.SubElement(cat, 'text')
        cat_text.text = path
        self.current_subcategory = name
        logger.info(f"Добавлена подкатегория: {name}")
    
    def _generate_permutations(self, answer: str, question_text: str = '') -> List[str]:
        """Генерирует перестановки для ответа.
        
        Перестановки генерируются ТОЛЬКО если:
        1. Ответ целиком состоит ТОЛЬКО из цифр (без пробелов, запятых, точек)
        2. Вопрос содержит ключевые слова из области РУССКОГО ЯЗЫКА:
           - цифры/цифр/номера (в контексте указания цифр)
           - по порядку/в порядке/по возрастанию/по убыванию
           - без пробелов/без пробела
        """
        # Проверяем, что ответ целиком состоит ТОЛЬКО из цифр
        answer_clean = answer.strip()
        if not answer_clean.isdigit():
            return [answer]
        
        digits = answer_clean
        
        if len(digits) < 3:
            return [answer]
        
        question_lower = question_text.lower() if question_text else ''
        
        # Ключевые слова, запрещающие перестановки (строгий порядок)
        no_permutation_keywords = [
            'в порядке возрастания', 'по возрастанию',
            'в порядке убывания', 'по убыванию',
        ]
        
        # Если требуется строгий порядок — возвращаем единственный ответ
        if any(kw in question_lower for kw in no_permutation_keywords):
            return [answer]
        
        # Ключевые слова, разрешающие перестановки (порядок не важен)
        permutation_keywords = [
            'цифр', 'номера', 'номеров',
            'по порядку', 'в порядке',
            'без пробелов', 'без пробела',
        ]
        
        needs_permutations = any(kw in question_lower for kw in permutation_keywords)
        
        if needs_permutations and len(digits) <= 7:
            return [''.join(p) for p in itertools.permutations(digits)]
        return [answer]


    def _generate_permutations_with_partial_scoring(self, all_answers_ordered: List[tuple], question_text: str) -> List[tuple]:
        """Генерирует все комбинации с partial scoring для ИСТ/ОБЩ.
        
        Генерирует ВСЕ перестановки от 1 до N (类似 оригиналу):
        - permutations('123456', 1) -> 1,2,3,4,5,6
        - permutations('123456', 2) -> 12,13,14,...21,23,24...
        - ... до permutations('123456', 6) -> 654321
        
        Логика partial scoring:
        - 100%: все правильные без неправильных
        - 50%: ≥50% правильных + ≤1 неправильный ИЛИ все правильные + 1 неправильный
        - 0%: остальные
        
        Args:
            all_answers_ordered: список (текст, True/False) - ответы в порядке вопроса
        
        Returns:
            List of (answer_text, fraction) tuples
        """
        # all_answers_ordered = [(текст, True/False), ...] в порядке вопроса
        # Номера: 1, 2, 3, ..., len(all_answers_ordered)
        
        total_count = len(all_answers_ordered)
        if total_count == 0:
            return []
        
        # Находим номера правильных ответов (1-based)
        correct_indices = []
        for i, (text, is_correct) in enumerate(all_answers_ordered, 1):
            if is_correct:
                correct_indices.append(i)
        
        correct_count = len(correct_indices)
        if correct_count == 0:
            return []
        
        results = []
        seen = set()
        
        # numbers = '123' для 3 ответов, '123456' для 6 ответов
        numbers = ''.join(str(i) for i in range(1, total_count + 1))
        
        # Генерируем ВСЕ перестановки (как в оригинале)
        for length in range(1, total_count + 1):
            for perm in itertools.permutations(numbers, length):
                perm_str = ''.join(perm)
                if perm_str in seen:
                    continue
                seen.add(perm_str)
                
                # Считаем сколько правильных и неправильных в комбинации
                correct_in_combo = sum(1 for p in perm_str if int(p) in correct_indices)
                wrong_in_combo = length - correct_in_combo
                
                # Определяем fraction
                if correct_in_combo == correct_count and wrong_in_combo == 0:
                    # Все правильные
                    fraction = 100
                elif correct_in_combo >= correct_count * 0.5 and wrong_in_combo <= 1:
                    # ≥50% правильных и не более 1 неправильного
                    fraction = 50
                elif correct_in_combo == correct_count and wrong_in_combo == 1:
                    # Все правильные + 1 неправильный
                    fraction = 50
                else:
                    fraction = 0
                
                # Добавляем только fraction = 100 или 50 (без fraction = 0)
                if fraction > 0:
                    results.append((perm_str, fraction))
        
        return results


    def create_shortanswer(self, name: str, content: List[str], grade: float = 1.0, subject: str = ''):
        """Создает вопрос с кратким ответом.
        
        Args:
            subject: код предмета ('ист', 'общ') для partial scoring логики
        """
        tree = etree.fromstring(SHORTANSWER_TEMPLATE)
        clean_name = name.replace('I:', '').replace('I ', '').strip() if name.startswith(('I:', 'I ')) else name
        tree.find('name').find('text').text = clean_name
        tree.find('defaultgrade').text = f'{grade:.7f}'
        
        full_question_text = name
        
        question_parts = []
        # Для partial scoring: сохраняем все ответы с их статусом и порядком
        all_answers_ordered = []  # [(текст, True/False), ...]
        
        for item in content:
            item = item.strip()
            if not item:
                continue
            if item.startswith('+:'):
                all_answers_ordered.append((item.replace('+:', '').strip(), True))
            elif item.startswith('-:'):
                all_answers_ordered.append((item.replace('-:', '').strip(), False))
            elif not item.startswith('-'):
                question_parts.append(item)
        
        question_text_parts = []
        for part in question_parts:
            cleaned = strip_service_markers(part)
            if cleaned:
                question_text_parts.append(cleaned)
        
        # Если ответы не найдены через +:, ищем ответ без префикса
        # Паттерн: после S: текста идёт ответ (слово в верхнем регистре или просто текст)
        correct_answers = [a[0] for a in all_answers_ordered if a[1]]
        if not correct_answers and len(question_text_parts) > 1:
            # Берём последний элемент как ответ (если это не S: текст)
            potential_answer = question_text_parts[-1]
            if potential_answer and not potential_answer.startswith('S:'):
                all_answers_ordered.append((potential_answer, True))
                question_text_parts = question_text_parts[:-1]
        
        # Сохраняем оригинальный текст для numcombo (без <br>)
        question_text_raw = '\n'.join(question_text_parts)
        
        question_text = '<br>'.join(question_text_parts)
        question_text_with_header = full_question_text + ' ' + question_text
        question_text = ImageProcessor.process_text(question_text, self.image_data)
        question_text = FormulaProcessor.tex_to_latex(question_text)
        
        qt_elem = tree.find('questiontext')
        if '_IMAGE_' in question_text:
            quest_data = get_image(question_text)
            qt_elem.find('text').text = etree.CDATA('<p>' + remove_service_markers(quest_data[0].replace('_IMAGE_', '')) + '</p>')
            for img_name, img_data in quest_data[1]:
                img_elem = etree.SubElement(qt_elem, 'file', name=img_name, path='/', encoding='base64')
                img_elem.text = img_data.replace('_IMAGE_', '')
        else:
            question_html = '<p>' + remove_service_markers(question_text) + '</p>'
            qt_elem.find('text').text = etree.CDATA(question_html)
        
        # Для partial scoring (маркер {numerical_partial} или ИСТ/ОБЩ fallback)
        correct_answers = [a[0] for a in all_answers_ordered if a[1]]
        wrong_answers = [a[0] for a in all_answers_ordered if not a[1]]
        
        # Используем partial scoring только для маркеров
        use_partial = subject in ('numerical_partial', 'numerical_partial', 'partial')
        use_numcombo = subject in ('numerical_numcombo', 'numerical_numcombo')
        
        # Для numerical_partial и numerical_numcombo используем шаблон numerical
        use_numerical_template = use_partial or use_numcombo
        
        if use_numerical_template:
            tree = etree.fromstring(NUMERICAL_TEMPLATE)
            # Удаляем placeholder answer из шаблона
            default_answer = tree.find('.//answer')
            if default_answer is not None:
                tree.remove(default_answer)
        #else:
        #    tree = etree.fromstring(SHORTANSWER_TEMPLATE)  # tree already exists from line 700
        
        # Получаем qt_elem из НОВОГО tree (после возможной замены шаблона)
        qt_elem = tree.find('questiontext')
        
        # Восстанавливаем name и grade (могут сброситься при замене tree)
        clean_name = name.replace('I:', '').replace('I ', '').strip() if name.startswith(('I:', 'I ')) else name
        tree.find('name').find('text').text = clean_name
        tree.find('defaultgrade').text = f'{grade:.7f}'
        
        # Проверяем numcombo РАНЬШЕ partial, чтобы маркер {numerical_numcombo} имел приоритет
        if use_numcombo:
            # {numerical_numcombo}: ответ из +: разбивается на цифры и генерируются все перестановки
            # Например: +: 123 -> ответы: 123, 132, 213, 231, 312, 321
            
            answer_text_raw = ''
            for item in content:
                item = item.strip()
                if item.startswith('+:'):
                    answer_text_raw = item.replace('+:', '').strip()
                    break
            
            answer_digits = [c for c in answer_text_raw if c.isdigit()]
            
            if len(answer_digits) == 1:
                ans_elem = etree.SubElement(tree, 'answer')
                ans_elem.set('fraction', '100')
                ans_elem.set('format', 'moodle_auto_format')
                etree.SubElement(ans_elem, 'text').text = answer_digits[0]
            elif len(answer_digits) > 1:
                for perm in itertools.permutations(answer_digits):
                    ans_elem = etree.SubElement(tree, 'answer')
                    ans_elem.set('fraction', '100')
                    ans_elem.set('format', 'moodle_auto_format')
                    etree.SubElement(ans_elem, 'text').text = ''.join(perm)
            
            question_text = ImageProcessor.process_text(question_text_raw, self.image_data)
            question_text = FormulaProcessor.tex_to_latex(question_text)
            question_text = question_text.replace('\n', '<br>')  # Сохраняем переносы строк
            
            qt_elem = tree.find('questiontext')
            if '_IMAGE_' in question_text:
                quest_data = get_image(question_text)
                qt_elem.find('text').text = etree.CDATA('<p>' + remove_service_markers(quest_data[0].replace('_IMAGE_', '')) + '</p>')
                for img_name, img_data in quest_data[1]:
                    img_elem = etree.SubElement(qt_elem, 'file', name=img_name, path='/', encoding='base64')
                    img_elem.text = img_data.replace('_IMAGE_', '')
            else:
                question_html = '<p>' + remove_service_markers(question_text) + '</p>'
                qt_elem.find('text').text = etree.CDATA(question_html)
            
            tree.find('.//correctfeedback/text').text = 'Ваш ответ верный.'
            tree.find('.//partiallycorrectfeedback/text').text = 'Ваш ответ частично правильный.'
            tree.find('.//incorrectfeedback/text').text = 'Ваш ответ неправильный.'
            
            self.root.append(tree)
            self.question_count += 1
            return
        elif use_partial:
            # {numerical_partial}: нумерация + partial scoring (100/50/0)
            # Создаём пронумерованный текст вопроса
            numbered_parts = []
            answer_positions = []  # [(позиция, текст, is_correct), ...]
            pos = 1
            
            for item in content:
                item = item.strip()
                if not item:
                    continue
                if item.startswith('+:'):
                    answer_positions.append((pos, item.replace('+:', '').strip(), True))
                    numbered_parts.append(f"{pos})" + item.replace('+:', '').strip())
                    pos += 1
                elif item.startswith('-:'):
                    answer_positions.append((pos, item.replace('-:', '').strip(), False))
                    numbered_parts.append(f"{pos})" + item.replace('-:', '').strip())
                    pos += 1
                elif not item.startswith('-'):
                    cleaned = strip_service_markers(item)
                    if cleaned:
                        numbered_parts.append(cleaned)
            
            # Обновляем question_text с нумерацией
            question_text = '<br>'.join(numbered_parts)
        question_text = ImageProcessor.process_text(question_text, self.image_data)
        question_text = FormulaProcessor.tex_to_latex(question_text)
        question_text = question_text.replace('\n', '<br>')  # Сохраняем переносы строк
        
        qt_elem = tree.find('questiontext')
        if '_IMAGE_' in question_text:
            quest_data = get_image(question_text)
            qt_elem.find('text').text = etree.CDATA('<p>' + remove_service_markers(quest_data[0].replace('_IMAGE_', '')) + '</p>')
            for img_name, img_data in quest_data[1]:
                img_elem = etree.SubElement(qt_elem, 'file', name=img_name, path='/', encoding='base64')
                img_elem.text = img_data.replace('_IMAGE_', '')
        else:
            question_html = '<p>' + remove_service_markers(question_text) + '</p>'
            qt_elem.find('text').text = etree.CDATA(question_html)
        
        # Генерируем ответы
        if use_partial and answer_positions:
            # numerical_partial: partial scoring
            answers_for_partial = [(text, is_correct) for pos, text, is_correct in answer_positions]
            partial_results = self._generate_permutations_with_partial_scoring(
                answers_for_partial, question_text_with_header
            )
            for ans_text, fraction in partial_results:
                ans_elem = etree.SubElement(tree, 'answer')
                ans_elem.set('fraction', str(fraction))
                ans_elem.set('format', 'moodle_auto_format')
                etree.SubElement(ans_elem, 'text').text = ans_text
        elif correct_answers:
            # Стандартная логика / {shortanswer_phrase} / fallback
            for answer_text in correct_answers:
                all_answers = self._generate_permutations(answer_text, question_text_with_header)
                for ans in all_answers:
                    ans_elem = etree.SubElement(tree, 'answer')
                    ans_elem.set('fraction', '100')
                    ans_elem.set('format', 'moodle_auto_format')
                    etree.SubElement(ans_elem, 'text').text = ans
                
                # Для одиночных ответов (без перестановок) — добавляем вариант с . / ,
                if len(all_answers) == 1:
                    ans = all_answers[0]
                    variant = ans.replace(',', '.') if ',' in ans else ans.replace('.', ',')
                    if variant != ans:
                        ans_elem2 = etree.SubElement(tree, 'answer')
                        ans_elem2.set('fraction', '100')
                        ans_elem2.set('format', 'moodle_auto_format')
                        etree.SubElement(ans_elem2, 'text').text = variant
        
        tree.find('.//correctfeedback/text').text = 'Ваш ответ верный.'
        tree.find('.//partiallycorrectfeedback/text').text = 'Ваш ответ частично правильный.'
        tree.find('.//incorrectfeedback/text').text = 'Ваш ответ неправильный.'
        
        self.root.append(tree)
        self.question_count += 1
        return
    
    def create_multichoice(self, name: str, content: List[str], grade: float = 1.0,
                           single: bool = None, penalty_wrong: int = 0):
        """Создает вопрос с выбором ответа.
        
        Args:
            single: True = один ответ, False = несколько. None = авто.
            penalty_wrong: штраф за неправильный (0 = без штрафа, -100 = полный штраф)
        """
        tree = etree.fromstring(MULTICHOICE_TEMPLATE)
        clean_name = name.replace('I:', '').replace('I ', '').strip() if name.startswith(('I:', 'I ')) else name
        tree.find('name').find('text').text = clean_name
        tree.find('defaultgrade').text = f'{grade:.7f}'
        
        question_parts = []
        answers = []
        right_count = 0
        
        for item in content:
            if not item.strip():
                continue
            
            # Проверяем, содержит ли строка несколько ответов через ;
            if ';' in item and (item.startswith('+:') or item.startswith('-:')):
                parsed = parse_answers_from_line(item)
                for ans_text, is_right in parsed:
                    if is_right:
                        right_count += 1
                    answers.append((ans_text, is_right))
            elif item.startswith('+:'):
                right_count += 1
                answers.append((item.replace('+:', '').strip(), True))
            elif item.startswith('+') and not item.startswith('+:'):
                right_count += 1
                answers.append((item[1:].strip(), True))
            elif item.startswith('-:'):
                answers.append((item.replace('-:', '').strip(), False))
            elif item.startswith('-'):
                answers.append((item[1:].strip(), False))
            else:
                question_parts.append(item)
        
        question_text_cleaned = []
        for part in question_parts:
            cleaned = strip_service_markers(part)
            if cleaned:
                question_text_cleaned.append(cleaned)
        
        question_text = '<br>'.join(question_text_cleaned)
        question_text = ImageProcessor.process_text(question_text, self.image_data)
        question_text = FormulaProcessor.tex_to_latex(question_text)
        
        qt_elem = tree.find('questiontext')
        if '_IMAGE_' in question_text:
            quest_data = get_image(question_text)
            qt_elem.find('text').text = etree.CDATA('<p>' + remove_service_markers(quest_data[0].replace('_IMAGE_', '')) + '</p>')
            for img_name, img_data in quest_data[1]:
                img_elem = etree.SubElement(qt_elem, 'file', name=img_name, path='/', encoding='base64')
                img_elem.text = img_data.replace('_IMAGE_', '')
        else:
            question_html = '<p>' + remove_service_markers(question_text) + '</p>'
            qt_elem.find('text').text = etree.CDATA(question_html)
        
        for ans_text, is_right in answers:
            # Вычисление fraction в зависимости от режима
            if is_right:
                fraction = str(round(100 / right_count, 5)) if right_count > 0 else '100'
            else:
                # penalty_wrong: 0 = без штрафа, -100 = полный штраф за каждый неправильный
                if penalty_wrong != 0:
                    fraction = str(penalty_wrong)
                else:
                    fraction = '0'
            ans_text = ImageProcessor.process_text(ans_text, self.image_data)
            ans_text = FormulaProcessor.tex_to_latex(ans_text)
            
            if '_IMAGE_' in ans_text:
                ans_data = get_image(ans_text)
                ans_elem = etree.SubElement(tree, 'answer', fraction=fraction, format='html')
                etree.SubElement(ans_elem, 'text').text = etree.CDATA('<p>' + ans_data[0].replace('_IMAGE_', '') + '</p>')
                if ans_data[1]:
                    img_elem = etree.SubElement(ans_elem, 'file', name=ans_data[1], path='/', encoding='base64')
                    img_elem.text = ans_data[2].replace('_IMAGE_', '')
            else:
                ans_elem = etree.SubElement(tree, 'answer', fraction=fraction, format='html')
                etree.SubElement(ans_elem, 'text').text = etree.CDATA(f'<p>{ans_text}</p>')
                etree.SubElement(etree.SubElement(ans_elem, 'feedback', format='html'), 'text').text = ''
        
        # Установка single/multiple
        if single is not None:
            tree.find('single').text = 'true' if single else 'false'
        elif right_count > 1:
            tree.find('single').text = 'false'
        
        if right_count > 1 or (single is not None and not single):
            tree.find('shownumcorrect').text = ''
        
        tree.find('.//correctfeedback/text').text = 'Ваш ответ верный.'
        tree.find('.//partiallycorrectfeedback/text').text = 'Ваш ответ частично правильный.'
        tree.find('.//incorrectfeedback/text').text = 'Ваш ответ неправильный.'
        
        self.root.append(tree)
        self.question_count += 1
    
    def create_cloze(self, name: str, content: List[str], grade: float = 1.0):
        """Создает вопрос типа cloze (пропущенные слова)."""
        tree = etree.fromstring(CLOZE_TEMPLATE)
        clean_name = name.replace('I:', '').replace('I ', '').strip() if name.startswith(('I:', 'I ')) else name
        tree.find('name').find('text').text = clean_name
        tree.find('defaultgrade').text = f'{grade:.7f}'
        
        question_text = ''
        for item in content:
            if item.startswith('+:'):
                answer = item.replace('+:', '').strip()
                question_text += f'<p>+:{answer}</p>'
            elif not item.startswith('-:'):
                question_text += f'<p>{FormulaProcessor.tex_to_latex(remove_service_markers(item))}</p>'
        
        qt_elem = tree.find('questiontext')
        qt_elem.find('text').text = etree.CDATA(question_text)
        
        self.root.append(tree)
        self.question_count += 1
    
    def create_matching(self, name: str, content: List[str], grade: float = 1.0):
        """Создает вопрос на сопоставление.
        
        Особенности:
        - Все L и R учитываются (не только уникальные)
        - Если R повторяется, все вхождения используются
        - Лишние R (больше чем L) становятся дистракторами
        """
        tree = etree.fromstring(MATCHING_TEMPLATE)
        clean_name = name.replace('I:', '').replace('I ', '').strip() if name.startswith(('I:', 'I ')) else name
        tree.find('name').find('text').text = clean_name
        tree.find('defaultgrade').text = f'{grade:.7f}'
        
        question_parts = []
        left_items = {}  # {номер: текст}
        right_items = {}  # {номер: [текст1, текст2, ...]} - список для повторений
        numbered_items = {}
        
        for item in content:
            if '->' in item:
                pass
            elif '=>' in item:
                pass
            elif re.match(r'^L\d+:', item):
                num = re.match(r'^L(\d+):', item).group(1)
                rest = re.sub(r'^L\d+:', '', item).strip()
                # Проверяем, есть ли R на той же строке
                r_match = re.search(r'\s+R(\d+):', rest)
                if r_match:
                    l_text = rest[:r_match.start()].strip()
                    r_num = r_match.group(1)
                    r_text = rest[r_match.end():].strip()
                    left_items[num] = l_text
                    if r_num not in right_items:
                        right_items[r_num] = []
                    right_items[r_num].append(r_text)
                else:
                    left_items[num] = rest
            elif re.match(r'^R\d+:', item):
                num = re.match(r'^R(\d+):', item).group(1)
                text = re.sub(r'^R\d+:', '', item).strip()
                if num not in right_items:
                    right_items[num] = []
                right_items[num].append(text)
            elif re.match(r'^\d+:', item):
                num = re.match(r'^(\d+):', item).group(1)
                numbered_items[num] = re.sub(r'^\d+:', '', item).strip()
            elif not item.startswith('+:') and not item.startswith('-:'):
                cleaned = strip_service_markers(item)
                if cleaned:
                    question_parts.append(cleaned)
        
        # Подсчитываем количество L и R
        num_left = len(left_items)
        # right_items = {номер: [текст1, текст2, ...]}
        
        question_text = '<br>'.join(question_parts)
        question_text = FormulaProcessor.tex_to_latex(question_text)
        
        qt_elem = tree.find('questiontext')
        if '_IMAGE_' in question_text:
            quest_data = get_image(question_text)
            qt_elem.find('text').text = etree.CDATA('<p>' + remove_service_markers(quest_data[0].replace('_IMAGE_', '')) + '</p>')
            for img_name, img_data in quest_data[1]:
                img_elem = etree.SubElement(qt_elem, 'file', name=img_name, path='/', encoding='base64')
                img_elem.text = img_data.replace('_IMAGE_', '')
        else:
            question_html = '<p>' + remove_service_markers(question_text) + '</p>'
            qt_elem.find('text').text = etree.CDATA(question_html)
        
        # Создаём пары: L1->R1[0], L2->R2[0], ...
        # Если R[i] имеет несколько вхождений, используем все по порядку
        left_nums = sorted(left_items.keys(), key=lambda x: int(x) if x.isdigit() else 0)
        
        pair_index = 0
        for ln in left_nums:
            l_text = left_items[ln]
            if ln in right_items:
                r_list = right_items[ln]
                # Берём первое неиспользованное вхождение R с этим номером
                r_text = r_list[0] if r_list else ''
                subq = etree.SubElement(tree, 'subquestion', format='html')
                etree.SubElement(subq, 'text').text = etree.CDATA(f'<p>{FormulaProcessor.tex_to_latex(l_text)}</p>')
                answer = etree.SubElement(subq, 'answer')
                etree.SubElement(answer, 'text').text = r_text
                pair_index += 1
        
        # Добавляем дистракторы: все оставшиеся R (которые не спарены)
        for rn, r_list in right_items.items():
            if len(r_list) > 1:
                # Все вхождения после первого становятся дистракторами
                for r_text in r_list[1:]:
                    subq = etree.SubElement(tree, 'subquestion', format='html')
                    etree.SubElement(subq, 'text').text = ''
                    answer = etree.SubElement(subq, 'answer')
                    etree.SubElement(answer, 'text').text = r_text
            elif rn not in left_items:
                # R без пары L - дистрактор
                for r_text in r_list:
                    subq = etree.SubElement(tree, 'subquestion', format='html')
                    etree.SubElement(subq, 'text').text = ''
                    answer = etree.SubElement(subq, 'answer')
                    etree.SubElement(answer, 'text').text = r_text
        
        # Нумерованные элементы: phrase -> number (fallback)
        for num in sorted(numbered_items.keys(), key=int):
            subq = etree.SubElement(tree, 'subquestion', format='html')
            etree.SubElement(subq, 'text').text = etree.CDATA(f'<p>{FormulaProcessor.tex_to_latex(numbered_items[num])}</p>')
            answer = etree.SubElement(subq, 'answer')
            etree.SubElement(answer, 'text').text = num
        
        self.root.append(tree)
        self.question_count += 1
    
    def create_ddmatch(self, name: str, content: List[str], grade: float = 1.0):
        """Создает вопрос типа ddmatch (выбор из выпадающих списков).
        
        Формат в docx:
          S:Инструкция
          L1: вариант1 (ответы)
          L2: вариант2
          ...
          R1: вопрос1
          R2: вопрос2
          
        В XML ddmatch:
          <subquestion><text>Вопрос</text><answer>Ответ</answer></subquestion>
        """
        tree = etree.fromstring(DDMATCH_TEMPLATE)
        clean_name = name.replace('I:', '').replace('I ', '').strip() if name.startswith(('I:', 'I ')) else name
        tree.find('name').find('text').text = clean_name
        tree.find('defaultgrade').text = f'{grade:.7f}'
        
        left_items = {}
        right_items = {}
        question_parts = []
        
        for item in content:
            if re.match(r'^L\d+:', item):
                num = re.match(r'^L(\d+):', item).group(1)
                left_items[num] = re.sub(r'^L\d+:', '', item).strip()
            elif re.match(r'^R\d+:', item):
                num = re.match(r'^R(\d+):', item).group(1)
                right_items[num] = re.sub(r'^R\d+:', '', item).strip()
            elif not item.startswith('+:') and not item.startswith('-:') and not item.startswith('S:'):
                cleaned = strip_service_markers(item)
                if cleaned:
                    question_parts.append(cleaned)
        
        question_text = '<br>'.join(question_parts)
        question_text = FormulaProcessor.tex_to_latex(question_text)
        
        qt_elem = tree.find('questiontext')
        if '_IMAGE_' in question_text:
            quest_data = get_image(question_text)
            qt_elem.find('text').text = etree.CDATA('<p>' + remove_service_markers(quest_data[0].replace('_IMAGE_', '')) + '</p>')
            for img_name, img_data in quest_data[1]:
                img_elem = etree.SubElement(qt_elem, 'file', name=img_name, path='/', encoding='base64')
                img_elem.text = img_data.replace('_IMAGE_', '')
        else:
            question_html = '<p>' + remove_service_markers(question_text) + '</p>'
            qt_elem.find('text').text = etree.CDATA(question_html)
        
        for num in sorted(right_items.keys(), key=lambda x: int(x) if x.isdigit() else 0):
            if num in left_items:
                subq = etree.SubElement(tree, 'subquestion')
                subq.set('format', 'html')
                etree.SubElement(subq, 'text').text = etree.CDATA(f'<p>{FormulaProcessor.tex_to_latex(right_items[num])}</p>')
                answer_elem = etree.SubElement(subq, 'answer')
                answer_elem.set('format', 'html')
                etree.SubElement(answer_elem, 'text').text = etree.CDATA(f'<p>{FormulaProcessor.tex_to_latex(left_items[num])}</p>')
        
        self.root.append(tree)
        self.question_count += 1
    
    def create_gapselect(self, name: str, content: List[str], grade: float = 1.0):
        """Создает вопрос типа gapselect (выбор слова для пропуска).
        
        Формат в docx:
          S:Инструкция (Вставьте пропущенные слова... или Выберите правильный вариант...)
          Текст с (1), (2), ... маркерами пропусков
          A) option1; B) option2; C) option3; D) option4 (группа 1) - ВСЕ В ОДНОЙ СТРОКЕ
          A) option1
          B) option2
          C) option3
          D) option4  - ИЛИ ПО ОДНОЙ СТРОКЕ
          +:ABCD... (ключ ответов)
        """
        tree = etree.fromstring(GAPSELECT_TEMPLATE)
        clean_name = name.replace('I:', '').replace('I ', '').strip() if name.startswith(('I:', 'I ')) else name
        tree.find('name').find('text').text = clean_name
        tree.find('defaultgrade').text = f'{grade:.7f}'
        
        intro_text = ''
        question_text_parts = []
        options = []
        answer_key = ''
        
        for item in content:
            item_stripped = item.strip()
            if not item_stripped:
                continue
            
            if item_stripped.startswith('S:'):
                intro_text = strip_service_markers(item_stripped)
            elif re.match(r'^ОТВЕТ:\s*[A-DА-Г]+$', item_stripped):
                answer_key = re.sub(r'^ОТВЕТ:\s*', '', item_stripped)
            elif re.match(r'^Ответ:\s*[A-DА-Г\s]+$', item_stripped):
                answer_key = re.sub(r'^Ответ:\s*', '', item_stripped).replace(' ', '')
            elif re.match(r'^Ответы\s*:\s*[A-DА-Г]+$', item_stripped):
                answer_key = re.sub(r'^Ответы\s*:\s*', '', item_stripped).replace(' ', '')
            elif re.match(r'^\+:\s*[A-DА-Г]+\s*$', item_stripped):
                answer_key = re.sub(r'^\+:\s*', '', item_stripped).replace(' ', '')
            elif re.match(r'^\+:[A-DА-Г]+$', item_stripped):
                answer_key = re.sub(r'^\+:', '', item_stripped)
            elif re.match(r'^ответы\s*:\s*[A-DА-Г]+$', item_stripped, re.IGNORECASE):
                answer_key = re.sub(r'^ответы\s*:\s*', '', item_stripped, re.IGNORECASE).replace(' ', '')
            elif re.match(r'^ответ:\s*[A-DА-Г\s]+$', item_stripped, re.IGNORECASE):
                answer_key = re.sub(r'^ответ:\s*', '', item_stripped, re.IGNORECASE).replace(' ', '')
            elif re.match(r'^[A-DА-Г]\)', item_stripped) or re.match(r'^(\t| |)?\d*\.?\s*[A-DА-Г]\)', item_stripped):
                all_opts = re.findall(r'([A-DА-Г])\)\s*([^;]+?)(?=\s*[A-DА-Г]\)|;|$)', item_stripped)
                if all_opts:
                    for letter, text in all_opts:
                        options.append((text.strip(), letter))
                else:
                    single_opt = re.match(r'^([A-D])\)\s*(.+)$', item_stripped)
                    if single_opt:
                        options.append((single_opt.group(2).strip(), single_opt.group(1)))
            elif not item_stripped.startswith('+:') and not item_stripped.startswith('-:'):
                question_text_parts.append(item_stripped)
        
        full_question_text = intro_text + ('<br>' if intro_text and question_text_parts else '') + '<br>'.join(question_text_parts)
        full_question_text = FormulaProcessor.tex_to_latex(full_question_text)
        
        alpha_map = {'A': 1, 'B': 2, 'C': 3, 'D': 4, 'А': 1, 'Б': 2, 'В': 3, 'Г': 4}
        
        class GapReplacer:
            def __init__(self):
                self.idx = 0
            def __call__(self, match):
                i = self.idx
                self.idx += 1
                correct_letter = answer_key[i] if i < len(answer_key) else 'A'
                correct_pos = alpha_map.get(correct_letter, 1)
                return f'[[{correct_pos + i * 4}]]'
        
        full_question_text = re.sub(r'\(\d+\)', GapReplacer(), full_question_text)
        
        qt_elem = tree.find('questiontext')
        qt_elem.find('text').text = etree.CDATA(
            f'<p>{full_question_text}</p>'
        )
        
        for i, (text, _) in enumerate(options):
            group = (i // 4) + 1
            so_elem = etree.SubElement(tree, 'selectoption')
            so_elem.set('format', 'html')
            so_text = etree.SubElement(so_elem, 'text')
            so_text.text = etree.CDATA(f' {FormulaProcessor.tex_to_latex(text)}')
            so_group = etree.SubElement(so_elem, 'group')
            so_group.text = str(group)
        
        tree.find('.//correctfeedback/text').text = 'Ваш ответ верный.'
        tree.find('.//partiallycorrectfeedback/text').text = 'Ваш ответ частично правильный.'
        tree.find('.//incorrectfeedback/text').text = 'Ваш ответ неправильный.'
        
        self.root.append(tree)
        self.question_count += 1
    
    def create_shortanswer_numerical(self, name: str, content: List[str], grade: float = 1.0):
        """Создает вопрос shortanswer для числового ответа.
        
        Маркер {numerical}: . и , считаются одним и тем же.
        Генерирует два варианта ответа: с точкой и с запятой.
        """
        tree = etree.fromstring(SHORTANSWER_TEMPLATE)
        clean_name = name.replace('I:', '').replace('I ', '').strip() if name.startswith(('I:', 'I ')) else name
        tree.find('name').find('text').text = clean_name
        tree.find('defaultgrade').text = f'{grade:.7f}'
        
        question_parts = []
        correct_answers = []
        
        for item in content:
            item = item.strip()
            if not item:
                continue
            if item.startswith('+:'):
                correct_answers.append(item.replace('+:', '').strip())
            elif item.startswith('S:'):
                question_parts.append(strip_service_markers(item))
            elif not item.startswith('-:'):
                question_parts.append(item)
        
        question_text = '<br>'.join(question_parts)
        question_text = ImageProcessor.process_text(question_text, self.image_data)
        question_text = FormulaProcessor.tex_to_latex(question_text)
        
        qt_elem = tree.find('questiontext')
        if '_IMAGE_' in question_text:
            quest_data = get_image(question_text)
            qt_elem.find('text').text = etree.CDATA('<p>' + remove_service_markers(quest_data[0].replace('_IMAGE_', '')) + '</p>')
            for img_name, img_data in quest_data[1]:
                img_elem = etree.SubElement(qt_elem, 'file', name=img_name, path='/', encoding='base64')
                img_elem.text = img_data.replace('_IMAGE_', '')
        else:
            qt_elem.find('text').text = etree.CDATA(f'<p>{remove_service_markers(question_text)}</p>')
        
        for ans in correct_answers:
            # Основной ответ
            ans_elem = etree.SubElement(tree, 'answer')
            ans_elem.set('fraction', '100')
            ans_elem.set('format', 'moodle_auto_format')
            etree.SubElement(ans_elem, 'text').text = ans
            
            # Вариант с заменой . ↔ ,
            variant = ans.replace(',', '.') if ',' in ans else ans.replace('.', ',')
            if variant != ans:
                ans_elem2 = etree.SubElement(tree, 'answer')
                ans_elem2.set('fraction', '100')
                ans_elem2.set('format', 'moodle_auto_format')
                etree.SubElement(ans_elem2, 'text').text = variant
        
        tree.find('.//correctfeedback/text').text = 'Ваш ответ верный.'
        tree.find('.//partiallycorrectfeedback/text').text = 'Ваш ответ частично правильный.'
        tree.find('.//incorrectfeedback/text').text = 'Ваш ответ неправильный.'
        
        self.root.append(tree)
        self.question_count += 1
    
    def create_numerical(self, name: str, content: List[str], grade: float = 1.0):
        """Создает вопрос типа numerical (числовой ответ) — fallback.
        
        Формат в docx:
          S:Текст вопроса
          +:12345 (правильный числовой ответ)
        """
        tree = etree.fromstring(NUMERICAL_TEMPLATE)
        clean_name = name.replace('I:', '').replace('I ', '').strip() if name.startswith(('I:', 'I ')) else name
        tree.find('name').find('text').text = clean_name
        tree.find('defaultgrade').text = f'{grade:.7f}'
        
        question_text = ''
        correct_answer = ''
        
        for item in content:
            if item.startswith('S:'):
                question_text = strip_service_markers(item)
            elif item.startswith('+:'):
                correct_answer = item.replace('+:', '').strip()
            elif not item.startswith('-:'):
                cleaned = strip_service_markers(item)
                if question_text:
                    question_text += ' ' + cleaned
                else:
                    question_text = cleaned
        
        qt_elem = tree.find('questiontext')
        qt_elem.find('text').text = etree.CDATA(f'<p>{remove_service_markers(question_text)}</p>')
        
        # Устанавливаем правильный ответ
        ans_elem = tree.find('answer')
        ans_elem.find('text').text = correct_answer
        
        self.root.append(tree)
        self.question_count += 1
    
    def save(self, output_path: str):
        """Сохраняет XML в файл."""
        tree = etree.ElementTree(self.root)
        with open(output_path, 'wb') as f:
            tree.write(f, encoding='utf-8', xml_declaration=True, method='xml')
        logger.info(f"XML сохранен: {output_path} (вопросов: {self.question_count})")


# ============================================================
# ОСНОВНОЙ КЛАСС КОНВЕРТЕРА
# ============================================================

class MoodleConverter:
    """Универсальный конвертер Word -> Moodle XML.
    
    Этапы конвертации:
      1. Извлечение изображений из docx
      2. Чтение текста через docxlatex
      3. Парсинг вопросов с поддержкой 7 форматов маркеров
      4. Определение типа вопроса и генерация XML
      5. Сохранение результата
    """
    
    def __init__(self, input_path: str, output_path: str, selected_indices: List[int] = None):
        self.input_path = input_path
        self.output_path = output_path
        self.image_processor = ImageProcessor()
        self.generator = XMLGenerator()
        self.errors: List[str] = []
        self.current_subject = ""
        self.selected_indices = selected_indices  # Индексы выбранных вопросов
        self._question_index = 0  # Счётчик вопросов для фильтрации
        self._marker_overrides = {}  # Переопределения маркеров из GUI
    
    def convert(self) -> bool:
        """Основная функция конвертации."""
        logger.info(f"Начало конвертации: {self.input_path}")
        
        try:
            # Определяем предмет по имени файла
            import os
            filename = os.path.basename(self.input_path).lower()
            if 'ист' in filename:
                self.current_subject = 'ист'
            elif 'общ' in filename:
                self.current_subject = 'общ'
            else:
                self.current_subject = ''
            
            # 1. Извлекаем изображения
            image_data = self.image_processor.extract_images(self.input_path)
            self.generator.set_image_data(image_data)
            
            # 2. Читаем контент из docx
            if docxlatex is None:
                logger.error("docxlatex не установлен")
                return False
            
            doc = docxlatex.Document(self.input_path)
            lines = doc.get_text().split('\n')
            
            # 3. Парсим вопросы
            question_content: List[str] = []
            last_question_name = ""
            last_grade = 1.0
            current_marker = ""  # текущий маркер типа вопроса
            self._question_index = 0  # счётчик вопросов для фильтрации
            
            for line in lines:
                line = unicodedata.normalize('NFC', line.strip())
                
                if not line:
                    continue
                
                # --- Обработка категорий и маркеров ---
                # Проверяем V1/V2 в ЛЮБОМ месте (не только в начале файла)
                if line.startswith('V1:'):
                    # Сохраняем предыдущий вопрос ДО смены категории
                    if question_content:
                        self._save_question(question_content, last_question_name, last_grade, current_marker)
                        question_content = []
                    self.generator.add_category(line.replace('V1:', '').strip())
                    continue
                
                # Формат: {marker}V2: описание
                marker_match = re.match(r'^\{(\w+)\}V2:(.*)', line)
                if marker_match:
                    # Сохраняем предыдущий вопрос ДО смены подкатегории
                    if question_content:
                        self._save_question(question_content, last_question_name, last_grade, current_marker)
                        question_content = []
                    current_marker = marker_match.group(1)
                    subcategory_name = marker_match.group(2).strip()
                    if subcategory_name:
                        self.generator.add_subcategory(subcategory_name)
                    logger.info(f"Маркер типа: {{{current_marker}}}")
                    continue
                
                if line.startswith('V2:'):
                    # Сохраняем предыдущий вопрос ДО смены подкатегории
                    if question_content:
                        self._save_question(question_content, last_question_name, last_grade, current_marker)
                        question_content = []
                    self.generator.add_subcategory(line.replace('V2:', '').strip())
                    continue
                
                # ================================================================
                # ОБРАБОТКА МАРКЕРОВ ВОПРОСОВ
                # ================================================================
                # Поддерживаем 7 форматов заголовков вопросов.
                # Каждый формат нормализуется к каноническому виду "I:Задание N. ..."
                # ================================================================
                
                # Формат 1: Стандартный "I:Задание N."
                is_question = line.startswith('I:')
                
                if not is_question:
                    # ---------------------------------------------------------------
                    # Формат 2: Двойной "I I:Задание N."
                    # Пример: "I I:Задание 535. Фомичева К.М., T40-30, b=3"
                    # Причина: артефакт копирования из Word (два символа "I")
                    # Исправление: убираем "I " (3 символа), оставляем "I:Задание..."
                    # ---------------------------------------------------------------
                    if line.startswith('I I:'):
                        line = 'I:' + line[4:]  # "I I:" → "I:" (skip "I I:" entirely)
                        is_question = True
                    
                    # ---------------------------------------------------------------
                    # Формат 3: "I Задание N." (пробел вместо двоеточия)
                    # Пример: "I Задание 535. Фомичева К.М., T40-30, b=3"
                    # Причина: повреждённый маркер (пропущено двоеточие)
                    # Исправление: заменяем "I " на "I:"
                    # ---------------------------------------------------------------
                    elif line.startswith('I ') and 'Задание' in line:
                        line = 'I:' + line[2:]  # "I " → "I:"
                        is_question = True
                    
                    # ---------------------------------------------------------------
                    # Формат 4: ":Задание N." (потерян символ "I")
                    # Пример: ":Задание 222. Даттуева Э.А., Привалова Л.М., T35-17, b=1"
                    # Причина: повреждение при экспорте из Word
                    # Исправление: добавляем "I" перед ":"
                    # ---------------------------------------------------------------
                    elif line.startswith(':Задание'):
                        line = 'I:' + line[1:]  # ":Задание" → "I:Задание"
                        is_question = True
                    
                    # ---------------------------------------------------------------
                    # Формат 5: "Задание N. Author, ТЗN-M, b=N" (без префикса I:)
                    # Пример: "Задание 1. Даттуева Э.А., Привалова Л.М., ТЗ1-1, b=1"
                    # Причина: автор документа не ставил "I:" перед первым вопросом
                    # Исправление: добавляем "I:" в начало
                    # ---------------------------------------------------------------
                    elif line.startswith('Задание ') and ', b=' in line:
                        if re.search(r'ТЗ\d+-\d+|Т\d+-\d+|T\d+-\d+', line):
                            line = 'I:' + line  # "Задание..." → "I:Задание..."
                            is_question = True
                    
                    # ---------------------------------------------------------------
                    # Формат 6: "Kn-=mЗадание N." (мусор в начале строки)
                    # Пример: "Kn-=mЗадание 17. Даттуева Э.А., Привалова Л.М., ТЗ1-17, b=1"
                    # Причина: артефакт нумерации при экспорте из Word
                    # Исправление: удаляем мусорный префикс, оставляем "Задание..."
                    # ---------------------------------------------------------------
                    elif re.match(r'^[A-Za-z0-9_\-=]+Задание', line) and ', b=' in line:
                        if re.search(r'ТЗ\d+-\d+|Т\d+-\d+|T\d+-\d+', line):
                            match = re.match(r'^([A-Za-z0-9_\-=]+)(Задание.*)', line)
                            if match:
                                line = 'I:' + match.group(2)  # "Kn-=mЗадание..." → "I:Задание..."
                                is_question = True
                    
                    # ---------------------------------------------------------------
                    # Формат 7: "Author И.О., ТЗN-M, b=N" (только автор и ТЗ)
                    # Пример: "Селиванова В.А., ТЗ1-22, b=1"
                    # Причина: часть вопросов в НЯ 10кл без слова "Задание"
                    # Исправление: извлекаем номер из ТЗN-M и формируем "I:Задание M. ..."
                    # ---------------------------------------------------------------
                    elif re.match(r'^[А-Яа-яёЁ]+\s+[А-Яа-яёЁ]\.[А-Яа-яёЁ]\.,', line) and ', b=' in line:
                        tz_match = re.search(r'ТЗ(\d+)-(\d+)', line)
                        if tz_match:
                            q_num = tz_match.group(2)  # Номер вопроса (например, "22")
                            line = f'I:Задание {q_num}. {line}'  # → "I:Задание 22. Селиванова..."
                            is_question = True
                
                # --- Если это строка-заголовок вопроса ---
                if is_question:
                    # Сохраняем предыдущий вопрос
                    if question_content:
                        self._save_question(question_content, last_question_name, last_grade, current_marker)
                        question_content = []
                    
                    # Увеличиваем счётчик вопросов для фильтрации
                    self._question_index += 1
                    
                    # Извлекаем имя и балл нового вопроса
                    last_question_name = line
                    grade_match = re.search(r'b=(\d+)', line)
                    last_grade = float(grade_match.group(1)) if grade_match else 1.0
                    continue
                
                # --- Строка контента вопроса ---
                is_answer = (
                    line.startswith('+:') or 
                    line.startswith('-:') or
                    (len(line) > 1 and line[0] in '+-' and line[1] in ' \t')
                )
                
                # Добавляем строку если это ответ, подсказка или не служебная строка
                if line.startswith(('S:', '->', '=>')) or is_answer or \
                   not line.startswith(('V1:', 'V2:', 'I:')):
                    question_content.append(line)
            
            # Сохраняем последний вопрос
            if question_content:
                self._save_question(question_content, last_question_name, last_grade, current_marker)
            
            # 4. Сохраняем результат
            self.generator.save(self.output_path)
            
            logger.info(f"Конвертация завершена успешно: {self.output_path}")
            return True
            
        except Exception as e:
            error_msg = f"Ошибка при конвертации: {e}"
            logger.error(error_msg)
            self.errors.append(error_msg)
            self._save_error_log()
            return False
    
    def _save_question(self, content: List[str], name: str, grade: float, marker: str = ''):
        """Сохраняет вопрос в зависимости от его типа."""
        logger.info(f"Попытка сохранить вопрос: {name[:50]}... (индекс: {self._question_index})")
        
        # Проверяем, нужно ли сохранить этот вопрос
        if self.selected_indices is not None:
            idx = self._question_index - 1
            logger.info(f"  Проверка фильтра: idx={idx}")
            if idx not in self.selected_indices:
                logger.info(f"  -> ПРОПУЩЕН (не в списке выбранных)")
                return  # Пропускаем невыбранный вопрос
            logger.info(f"  -> ВЫБРАН")
        else:
            logger.info(f"  -> Нет фильтрации")
        
        # Используем маркер из _marker_overrides если установлен из GUI
        if name in self._marker_overrides:
            marker = self._marker_overrides[name]
        
        try:
            q_type = QuestionTypeDetector.detect(content, self.current_subject, marker)
            logger.info(f"  Определён тип: {q_type}")
            
            if q_type == 'multichoice_one':
                self.generator.create_multichoice(name, content, grade, single=True, penalty_wrong=0)
            elif q_type == 'multichoice_many':
                self.generator.create_multichoice(name, content, grade, single=False, penalty_wrong=-100)
            elif q_type == 'shortanswer_phrase':
                self.generator.create_shortanswer(name, content, grade, subject='')
            elif q_type == 'numerical_partial':
                self.generator.create_shortanswer(name, content, grade, subject='numerical_partial')
            elif q_type == 'numerical_numcombo':
                logger.info(f"  -> вызов create_shortanswer с numerical_numcombo")
                self.generator.create_shortanswer(name, content, grade, subject='numerical_numcombo')
            elif q_type == 'numerical':
                self.generator.create_shortanswer_numerical(name, content, grade)
            elif q_type == 'matching':
                self.generator.create_matching(name, content, grade)
            elif q_type == 'match_123':
                self.generator.create_matching(name, content, grade)
            elif q_type == 'ddmatch':
                self.generator.create_ddmatch(name, content, grade)
            elif q_type == 'gapselect':
                self.generator.create_gapselect(name, content, grade)
            elif q_type == 'cloze':
                self.generator.create_cloze(name, content, grade)
            # --- Fallback (старая эвристика) ---
            elif q_type == 'shortanswer':
                self.generator.create_shortanswer(name, content, grade, self.current_subject)
            elif q_type == 'multichoice':
                self.generator.create_multichoice(name, content, grade)
            else:
                self.generator.create_shortanswer(name, content, grade)
                
            logger.info(f"  -> Вопрос добавлен")
            
        except Exception as e:
            error_msg = f"Ошибка при сохранении вопроса '{name}': {e}"
            logger.error(error_msg)
            self.errors.append(error_msg)
    
    def _save_error_log(self):
        """Сохраняет лог ошибок."""
        if self.errors:
            error_log_path = os.path.join(
                LOG_DIR, 
                f"errors_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
            )
            with open(error_log_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(self.errors))
            logger.info(f"Лог ошибок сохранен: {error_log_path}")


# ============================================================
# ФУНКЦИИ ДЛЯ ПАКЕТНОЙ ОБРАБОТКИ
# ============================================================

def convert_folder(input_folder: str, output_folder: str):
    """Конвертирует все docx файлы из папки."""
    os.makedirs(output_folder, exist_ok=True)
    
    docx_files = []
    for root, dirs, files in os.walk(input_folder):
        for file in files:
            if file.endswith('.docx'):
                docx_files.append(os.path.join(root, file))
    
    logger.info(f"Найдено файлов для конвертации: {len(docx_files)}")
    
    success_count = 0
    error_count = 0
    
    for docx_path in docx_files:
        try:
            output_path = os.path.join(
                output_folder,
                os.path.splitext(os.path.basename(docx_path))[0] + '.xml'
            )
            
            converter = MoodleConverter(docx_path, output_path)
            if converter.convert():
                success_count += 1
            else:
                error_count += 1
                
        except Exception as e:
            logger.error(f"Ошибка при обработке {docx_path}: {e}")
            error_count += 1
    
    logger.info(f"Конвертация завершена. Успешно: {success_count}, Ошибок: {error_count}")


# ============================================================
# ТОЧКА ВХОДА
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description='Универсальный конвертер Word -> Moodle XML'
    )
    parser.add_argument('input', help='Путь к файлу Word или папке')
    parser.add_argument('--output', '-o', help='Путь для сохранения XML (если input файл)')
    parser.add_argument('--output-folder', help='Папка для сохранения XML (если input папка)')
    
    args = parser.parse_args()
    
    if os.path.isdir(args.input):
        # Папка — обрабатываем все файлы
        output_folder = args.output_folder or args.input.replace('input', 'output')
        convert_folder(args.input, output_folder)
    else:
        # Один файл
        if not args.output:
            base_name = os.path.splitext(os.path.basename(args.input))[0]
            args.output = f"{base_name}_moodle.xml"
        
        converter = MoodleConverter(args.input, args.output)
        converter.convert()


if __name__ == '__main__':
    main()
